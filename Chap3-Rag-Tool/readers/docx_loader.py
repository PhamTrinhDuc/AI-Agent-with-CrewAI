import sys
import docx
import unicodedata
import pandas as pd
from pathlib import Path
from typing import List, Optional
from llama_index.core.readers.base import BaseReader
from llama_index.core import Document
from utils import split_text



class DocxReader(BaseReader):
    """Read Docx files that respect table, using python-docx library

    Reader behavior:
        - All paragraph are extracted as a Document
        - Each table is extracted as a Document, rendered as a CSV string
        - The output is a list of Documents, concatenating the above
        (tables + paragraphs)
    """

    def __init__(self, max_words_per_page: int = 2048, *args, **kwargs):
        self.max_words_per_page = max_words_per_page

        try:
            import docx
        except ImportError:
            raise ImportError(
                "docx is not installed."
                "Please install it using `pip install python-docx`"
            )
        
    def _load_single_table(self, table) -> list[List[str]]:
        """
        Extract content from tables. Return a list of columns: list[str]
        Some merged cells will share duplicated content.
        """
        n_row = len(table.rows)
        n_col = len(table.columns)

        arrays = [["" for j in range(n_col)] for i in range(n_row)]

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                arrays[i][j] = cell.text
        
        return arrays
    
    def load_data(self, 
                  file_path: Path, 
                  extra_info: Optional[dict] = None, 
                  **kwargs) -> List[Document]:
        """Load data using Docx reader

        Args:
            file_path (Path): Path to .docx file

        Returns:
            List[Document]: list of documents extracted from the HTML file
        """

        file_path = Path(file_path).resolve()
        doc = docx.Document(docx = str(file_path))
        all_text = "\n".join(
            [unicodedata.normalize("NFKC", p.text) for p in doc.paragraphs]
        )

        pages = split_text(text=all_text, max_tokens=self.max_words_per_page)
        
        # convert table in document to dataframe
        tables = []
        for table in doc.tables:
            arrays = self._load_single_table(table=table)
            tables.append(pd.DataFrame({a[0]: a[1:] for a in arrays}))
        
        extra_info = extra_info or {}
        # create output Document with metadata from table
        documents = [
            Document(
                text=table.to_csv(index = False).strip(),
                metadata = {
                    "table_origin": table.to_csv(index = False),
                    "type": "table",
                    **extra_info
                },
                metadata_seperator="",
                metadata_template=""
            )
            for table in tables
        ]
        # create Document from non-table text
        documents.extend(
            [
                Document(
                    text=non_table_text.strip(),
                    metadata = {
                        "page_label": i + 1, 
                        **extra_info
                    }
                )
                for i, non_table_text in enumerate(pages)
            ]
        )
        return documents

    def test(self, data_path: str = None):
        path = data_path or "Chap3-Rag-Tool/data/demo/demo.docx"
        doc = docx.Document(docx=path)
        
        # see table in document
        # for i, table in enumerate(doc.tables):
        #     print(f"Table {i+1}" + "=" * 50)
        #     for row in table.rows:
        #         for cell in row.cells:
        #             print(cell.text, end= " | ")
        #         print()
        
        # see paragraph in document
        # for i, p in enumerate(doc.paragraphs):
        #     print(f"Paragraph {i+1}: {p.text}")

        # test method _load_single_table
        # arrays = self._l
        
        # test method load_data
        documents =self.load_data(file_path=path)
        print(documents[-1])

        